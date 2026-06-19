"""
Agent 3: Soạn thảo Thông báo kết luận cuộc họp
Input: PDF/DOCX/TXT (giấy mời, biên bản, công văn)
Output: File Word Thông báo kết luận theo thể thức hành chính
"""
import os
import sys
import re
import glob as gb
from datetime import datetime

sys.stdout.reconfigure(encoding='utf-8')

# Templates
HEADER = """ỦY BAN NHÂN DÂN
PHƯỜNG BÀ RỊA
─────────────
Số:    /TB-UBND

CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM
Độc lập – Tự do – Hạnh phúc
──────────────────────────
Bà Rịa, ngày {ngay} tháng {thang} năm {nam}

THÔNG BÁO KẾT LUẬN
{ten_cuoc_hop}

I. THÔNG TIN CUỘC HỌP
- Thời gian: {thoi_gian}
- Địa điểm: {dia_diem}
- Chủ trì: {chu_tri}
- Thành phần tham dự: {thanh_phan}

II. NỘI DUNG BÁO CÁO
{noi_dung_bao_cao}

III. Ý KIẾN CHỈ ĐẠO VÀ KẾT LUẬN
{y_kien_chi_dao}

IV. NHIỆM VỤ ĐƯỢC GIAO
{nhiem_vu_duoc_giao}

Văn phòng UBND phường Bà Rịa thông báo để các đơn vị liên quan biết, thực hiện./.

Nơi nhận:
- {noi_nhan};
- Lưu: VT.

TM. ỦY BAN NHÂN DÂN
CHỦ TỊCH"""


def load_credentials():
    google_auth_available = False
    try:
        from google.oauth2.credentials import Credentials
        from googleapiclient.discovery import build
        google_auth_available = True
    except ImportError:
        pass

    try:
        from dotenv import load_dotenv
        load_dotenv()
    except ImportError:
        pass

    return google_auth_available


def read_input(path):
    text = ""
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text += t + "\n"
        except ImportError:
            print("  ❌ Cần cài: pip install pdfplumber")
            return ""
    elif ext == ".docx":
        try:
            from docx import Document
            doc = Document(path)
            for p in doc.paragraphs:
                text += p.text + "\n"
        except ImportError:
            print("  ❌ Cần cài: pip install python-docx")
            return ""
    elif ext == ".txt":
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
    else:
        print(f"  ❌ Không hỗ trợ định dạng: {ext}")
        return ""

    return text.strip()


def parse_meeting_info(text, filename=""):
    info = {
        "ten_cuoc_hop": "",
        "thoi_gian": "",
        "dia_diem": "",
        "chu_tri": "",
        "thanh_phan": "",
        "noi_dung_bao_cao": "",
        "y_kien_chi_dao": "",
        "nhiem_vu_duoc_giao": "",
        "noi_nhan": "Như trên",
    }

    lines = text.split("\n")

    # Tên cuộc họp
    for line in lines:
        ls = line.strip()
        if ls and not any(kw in ls.upper() for kw in ["ỦY BAN", "CỘNG HÒA", "ĐỘC LẬP", "SỐ:", "───"]):
            if len(ls) > 10:
                info["ten_cuoc_hop"] = ls
                break

    if not info["ten_cuoc_hop"]:
        info["ten_cuoc_hop"] = f"(Cuộc họp theo văn bản: {filename})"

    # Thời gian
    for line in lines:
        if any(kw in line.lower() for kw in ["thời gian", "vào lúc", "lúc "]):
            info["thoi_gian"] = re.sub(r"^[-•\s]*(Thời gian|thời gian)\s*[:\-]?\s*", "", line).strip()
            break

    # Địa điểm
    for line in lines:
        if "địa điểm" in line.lower() or line.strip().startswith("Địa"):
            info["dia_diem"] = re.sub(r"^[-•\s]*(Địa điểm|địa điểm|Địa)\s*[:\-]?\s*", "", line).strip()
            break

    # Chủ trì
    for line in lines:
        if "chủ trì" in line.lower():
            info["chu_tri"] = re.sub(r"^[-•\s]*(Chủ trì|chủ trì)\s*[:\-]?\s*", "", line).strip()
            break

    # Thành phần
    tp_parts = []
    in_tp = False
    for line in lines:
        ls = line.strip()
        if any(kw in ls.lower() for kw in ["kính mời", "tham dự", "thành phần"]):
            in_tp = True
            continue
        if in_tp:
            if ls.startswith("-") or ls.startswith("•"):
                tp_parts.append(ls.lstrip("-• ").strip())
            elif any(kw in ls.lower() for kw in ["chủ trì", "địa điểm", "thời gian", "nơi nhận", "trân trọng"]):
                break
    if tp_parts:
        info["thanh_phan"] = "\n".join(f"- {p}" for p in tp_parts)

    # Nội dung báo cáo
    nd_lines = []
    in_nd = False
    for line in lines:
        ls = line.strip()
        if any(kw in ls.lower() for kw in ["về nội dung", "về việc", "nội dung"]):
            in_nd = True
        if in_nd:
            if ls and not any(kw in ls.upper() for kw in ["CHỦ TRÌ", "KẾT LUẬN", "NHIỆM VỤ", "NƠI NHẬN"]):
                nd_lines.append(ls)
            elif any(kw in ls.upper() for kw in ["KẾT LUẬN", "NHIỆM VỤ"]):
                break
    if nd_lines:
        info["noi_dung_bao_cao"] = "\n".join(nd_lines)

    # Ý kiến chỉ đạo (từ phần còn lại)
    if not info["y_kien_chi_dao"]:
        info["y_kien_chi_dao"] = "Chủ trì cuộc họp ghi nhận các ý kiến tham gia và kết luận các nội dung cụ thể."

    # Nhiệm vụ được giao
    nv_lines = []
    in_nv = False
    for line in lines:
        ls = line.strip()
        if "nhiệm vụ" in ls.lower():
            in_nv = True
            continue
        if in_nv:
            if ls and not any(kw in ls.upper() for kw in ["NƠI NHẬN", "CHỦ TỊCH", "LƯU"]):
                nv_lines.append(ls)
            else:
                break
    if nv_lines:
        info["nhiem_vu_duoc_giao"] = "\n".join(nv_lines)

    # Nơi nhận
    for line in lines:
        m = re.search(r"Nơi nhận\s*:\s*(.+)", line, re.I)
        if m:
            info["noi_nhan"] = m.group(1).strip()
            break

    return info


def create_word(info, output_path):
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # Fallback to TXT
        content = HEADER.format(**info)
        with open(output_path.replace(".docx", ".txt"), "w", encoding="utf-8") as f:
            f.write(content)
        print(f"  ✅ Đã lưu (txt): {output_path.replace('.docx', '.txt')}")
        return output_path.replace(".docx", ".txt")

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    font = doc.styles['Normal'].font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    def add_line(text, bold=False, size=13, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(0), space_before=Pt(0)):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = space_after
        p.paragraph_format.space_before = space_before
        p.paragraph_format.line_spacing = 1.3
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold
        return p

    add_line("ỦY BAN NHÂN DÂN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    add_line("PHƯỜNG BÀ RỊA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    add_line("─────────────", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_line(f"Số:    /TB-UBND", align=WD_ALIGN_PARAGRAPH.RIGHT)

    add_line("")
    add_line("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    add_line("Độc lập – Tự do – Hạnh phúc", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    add_line("──────────────────────────", align=WD_ALIGN_PARAGRAPH.CENTER)

    now = datetime.now()
    add_line(f"Bà Rịa, ngày {now.day} tháng {now.month} năm {now.year}", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_line("")
    add_line("THÔNG BÁO KẾT LUẬN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_line(info["ten_cuoc_hop"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)

    add_line("")
    add_line("I. THÔNG TIN CUỘC HỌP", bold=True, size=13, space_before=Pt(6))
    add_line(f"- Thời gian: {info['thoi_gian']}")
    add_line(f"- Địa điểm: {info['dia_diem']}")
    add_line(f"- Chủ trì: {info['chu_tri']}")
    add_line(f"- Thành phần tham dự: {info['thanh_phan']}")

    add_line("")
    add_line("II. NỘI DUNG BÁO CÁO", bold=True, size=13, space_before=Pt(6))
    for line in info['noi_dung_bao_cao'].split("\n"):
        add_line(line)

    add_line("")
    add_line("III. Ý KIẾN CHỈ ĐẠO VÀ KẾT LUẬN", bold=True, size=13, space_before=Pt(6))
    for line in info['y_kien_chi_dao'].split("\n"):
        add_line(line)

    add_line("")
    add_line("IV. NHIỆM VỤ ĐƯỢC GIAO", bold=True, size=13, space_before=Pt(6))
    for line in info['nhiem_vu_duoc_giao'].split("\n"):
        add_line(line)

    add_line("")
    add_line("Văn phòng UBND phường Bà Rịa thông báo để các đơn vị liên quan biết, thực hiện./.")
    add_line("")

    add_line("Nơi nhận:")
    add_line(f"- {info['noi_nhan']};")
    add_line("- Lưu: VT.")

    add_line("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run("TM. ỦY BAN NHÂN DÂN\nCHỦ TỊCH")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = True

    doc.save(output_path)
    return output_path


def run(input_path=None):
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    os.chdir(base_dir)

    print("=" * 55)
    print("  📝 Agent 3: Thông báo kết luận cuộc họp")
    print("=" * 55)

    load_credentials()

    if not input_path:
        pdfs = gb.glob("scripts/*.pdf") + gb.glob("*.pdf") + gb.glob("agents/*.pdf")
        if pdfs:
            print(f"\n📄 Tìm thấy {len(pdfs)} file PDF:")
            for i, p in enumerate(pdfs, 1):
                print(f"   {i}. {p}")
            choice = input(f"\nChọn file (1-{len(pdfs)}, Enter = file đầu): ").strip()
            if choice.isdigit() and 1 <= int(choice) <= len(pdfs):
                input_path = pdfs[int(choice) - 1]
            else:
                input_path = pdfs[0]
        else:
            input_path = input("\nNhập đường dẫn file (PDF/DOCX/TXT): ").strip()
            if not input_path:
                print("⚠️ Không có file đầu vào")
                return

    print(f"\n📂 Đọc file: {input_path}")
    text = read_input(input_path)
    if not text:
        print("❌ Không đọc được nội dung")
        return

    print("  ✅ Đã đọc nội dung")
    print(f"     {len(text)} ký tự")

    info = parse_meeting_info(text, os.path.basename(input_path))

    print("\n📋 Thông tin trích xuất:")
    print(f"   📌 Tên cuộc họp: {info['ten_cuoc_hop'][:60]}")
    print(f"   🕐 Thời gian: {info['thoi_gian'][:50]}")
    print(f"   📍 Địa điểm: {info['dia_diem'][:50]}")
    print(f"   👤 Chủ trì: {info['chu_tri'][:50]}")

    out_dir = os.path.join(base_dir, "output", "ket_luan")
    os.makedirs(out_dir, exist_ok=True)
    now = datetime.now()
    out_name = f"Thong_bao_ket_luan_{now.strftime('%Y%m%d_%H%M%S')}.docx"
    out_path = os.path.join(out_dir, out_name)

    print(f"\n📄 Đang tạo văn bản...")
    result = create_word(info, out_path)

    print(f"\n✅ Đã lưu: {result}")
    return result


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Agent soạn thảo Thông báo kết luận")
    parser.add_argument("--input", "-i", help="File biên bản/giấy mờ/i (.pdf, .docx, .txt)")
    parser.add_argument("--yes", "-y", action="store_true", help="Tự động chọn file đầu tiên")
    args = parser.parse_args()
    run(args.input)


if __name__ == "__main__":
    main()
