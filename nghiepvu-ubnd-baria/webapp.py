#!/usr/bin/env python3
"""
Web App - UBND Phường Bà Rịa
Flask web interface cho hệ thống tự động hóa nghiệp vụ hành chính
"""
import os
import sys
import glob
import re
from functools import wraps
from datetime import datetime, date, timezone, timedelta
from io import BytesIO

sys.stdout.reconfigure(encoding='utf-8')
BASE = os.path.dirname(os.path.abspath(__file__))
SCRIPTS = os.path.join(BASE, "scripts")
DATA_DIR = ""

from dotenv import load_dotenv
load_dotenv(dotenv_path=os.path.join(SCRIPTS, ".env"), override=True)

_dd = os.getenv("DATA_DIR", "")
if _dd:
    DATA_DIR = _dd
    SCRIPTS = os.path.join(_dd, "scripts")
    os.makedirs(os.path.join(_dd, "instance"), exist_ok=True)
    os.makedirs(SCRIPTS, exist_ok=True)
    os.makedirs(os.path.join(_dd, "config"), exist_ok=True)
    os.makedirs(os.path.join(_dd, "output"), exist_ok=True)
    os.makedirs(os.path.join(_dd, "logs"), exist_ok=True)

from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory, session
from models import db, User
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build

TOKEN = os.getenv("TOKEN_PATH", "token.json")
LICH_SHEET_ID = os.getenv("GOOGLE_SHEET_LICH_ID", "")
TASK_SHEET_ID = os.getenv("GOOGLE_SHEET_TASK_ID", "")
CALENDAR_ID = os.getenv("GOOGLE_CALENDAR_ID", "primary")
SHEET_SCOPES = ["https://www.googleapis.com/auth/spreadsheets"]
CAL_SCOPES = ["https://www.googleapis.com/auth/calendar"]

token_path = os.path.join(DATA_DIR, "config", TOKEN) if DATA_DIR else os.path.join(SCRIPTS, TOKEN)


def get_creds(scopes):
    return Credentials.from_authorized_user_file(token_path, scopes)


def sheets_svc():
    return build('sheets', 'v4', credentials=get_creds(SHEET_SCOPES))


def cal_svc():
    return build('calendar', 'v3', credentials=get_creds(CAL_SCOPES))


def read_sheet(sheet_id, range_):
    r = sheets_svc().spreadsheets().values().get(spreadsheetId=sheet_id, range=range_).execute()
    return r.get('values', [])


def append_sheet(sheet_id, range_, values):
    sheets_svc().spreadsheets().values().append(
        spreadsheetId=sheet_id, range=range_,
        valueInputOption='USER_ENTERED',
        body={'values': [values]}
    ).execute()


def update_sheet(sheet_id, range_, values):
    sheets_svc().spreadsheets().values().update(
        spreadsheetId=sheet_id, range=range_,
        valueInputOption='USER_ENTERED',
        body={'values': [values]}
    ).execute()


def delete_row(sheet_id, row_index):
    m = sheets_svc().spreadsheets().get(spreadsheetId=sheet_id).execute()
    sid = m['sheets'][0]['properties']['sheetId']
    sheets_svc().spreadsheets().batchUpdate(spreadsheetId=sheet_id, body={
        'requests': [{'deleteDimension': {
            'range': {'sheetId': sid, 'dimension': 'ROWS',
                      'startIndex': row_index - 1, 'endIndex': row_index}
        }}]
    }).execute()


app = Flask(__name__, template_folder=os.path.join(BASE, 'templates'))
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'change-me-in-production')
app.config['SESSION_COOKIE_NAME'] = 'ubnd_session'
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv("DATABASE_URL", f"sqlite:///{DATA_DIR}/instance/ubnd.db" if DATA_DIR else "sqlite:///ubnd.db")
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db.init_app(app)

with app.app_context():
    db.create_all()


def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user' not in session:
            return redirect('/login')
        return f(*args, **kwargs)
    return decorated


@app.context_processor
def inject_globals():
    return {
        'task_sheet_id': TASK_SHEET_ID,
        'lich_sheet_id': LICH_SHEET_ID,
        'now': datetime.now().strftime('%d/%m/%Y %H:%M'),
        'current_user': session.get('user'),
    }

# ===================== AUTH & LANDING =====================

@app.route('/')
def landing():
    if 'user' in session:
        return redirect('/dashboard')
    return render_template('landing.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'user' in session:
        return redirect('/dashboard')
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        user = User.query.filter_by(email=email).first()
        if user and user.check_password(password):
            session['user'] = {'email': user.email, 'fullname': user.fullname}
            flash('Đăng nhập thành công!', 'success')
            return redirect('/dashboard')
        return render_template('login.html', error='Email hoặc mật khẩu không đúng.')
    return render_template('login.html')


@app.route('/register', methods=['GET', 'POST'])
def register():
    if 'user' in session:
        return redirect('/dashboard')
    if request.method == 'POST':
        fullname = request.form.get('fullname', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        confirm = request.form.get('confirm_password', '')
        if not fullname or not email or not password:
            return render_template('register.html', error='Vui lòng điền đầy đủ thông tin.')
        if password != confirm:
            return render_template('register.html', error='Mật khẩu xác nhận không khớp.')
        if len(password) < 6:
            return render_template('register.html', error='Mật khẩu tối thiểu 6 ký tự.')
        if User.query.filter_by(email=email).first():
            return render_template('register.html', error='Email đã được đăng ký.')
        user = User(email=email, fullname=fullname)
        user.set_password(password)
        db.session.add(user)
        db.session.commit()
        session['user'] = {'email': email, 'fullname': fullname}
        flash('Đăng ký thành công!', 'success')
        return redirect('/dashboard')
    return render_template('register.html')


@app.route('/logout')
def logout():
    session.pop('user', None)
    flash('Đã đăng xuất.', 'info')
    return redirect('/')


# ===================== DASHBOARD =====================

@app.route('/dashboard')
@login_required
def dashboard():
    rows = read_sheet(TASK_SHEET_ID, 'A:L')
    tasks = rows[1:] if len(rows) > 1 else []
    task_count = len(tasks)
    done = sum(1 for t in tasks if len(t) > 9 and 'Hoàn thành' in t[9])
    overdue = sum(1 for t in tasks if len(t) > 9 and 'Quá hạn' in t[9])
    soon = []
    for t in tasks:
        if len(t) > 4 and t[4]:
            try:
                p = t[4].strip().split('/')
                d = date(int(p[2]), int(p[1]), int(p[0]))
                diff = (d - date.today()).days
                if 0 <= diff <= 3:
                    soon.append((t[4], t[1] if len(t) > 1 else '', diff))
            except:
                pass
    soon.sort(key=lambda x: x[2])

    ev_rows = read_sheet(LICH_SHEET_ID, 'A:I')
    event_count = max(0, len(ev_rows) - 2)

    return render_template('index.html', active='dashboard', stats={
        'task_count': task_count, 'done': done, 'overdue': overdue,
        'soon': soon, 'event_count': event_count,
    })

# ===================== TASKS =====================

@app.route('/tasks')
@login_required
def tasks():
    rows = read_sheet(TASK_SHEET_ID, 'A:L')
    return render_template('tasks.html', active='tasks', tasks=rows[1:] if len(rows) > 1 else [])


@app.route('/tasks/add', methods=['GET', 'POST'])
@login_required
def task_add():
    if request.method == 'POST':
        row = [
            "",
            request.form.get('content', ''),
            request.form.get('so_vb', ''),
            request.form.get('ngay_bh', ''),
            request.form.get('han_xu_ly', ''),
            request.form.get('chu_tri', ''),
            request.form.get('phoi_hop', ''),
            request.form.get('nguoi_phu_trach', ''),
            request.form.get('nguoi_theo_doi', ''),
            request.form.get('trang_thai', '🟢 Đang thực hiện'),
            "", ""
        ]
        append_sheet(TASK_SHEET_ID, 'A:L', row)
        flash('Đã thêm nhiệm vụ!', 'success')
        return redirect('/tasks')
    return render_template('task_form.html', title='Thêm nhiệm vụ', icon='bi-plus-circle', is_edit=False, task=None)


@app.route('/tasks/edit/<int:idx>', methods=['GET', 'POST'])
@login_required
def task_edit(idx):
    rows = read_sheet(TASK_SHEET_ID, 'A:L')
    if idx >= len(rows):
        flash('Không tìm thấy nhiệm vụ', 'danger')
        return redirect('/tasks')
    t = rows[idx]
    if request.method == 'POST':
        new = [
            t[0] if len(t) > 0 else "",
            request.form.get('content', t[1] if len(t) > 1 else ''),
            request.form.get('so_vb', t[2] if len(t) > 2 else ''),
            request.form.get('ngay_bh', t[3] if len(t) > 3 else ''),
            request.form.get('han_xu_ly', t[4] if len(t) > 4 else ''),
            request.form.get('chu_tri', t[5] if len(t) > 5 else ''),
            request.form.get('phoi_hop', t[6] if len(t) > 6 else ''),
            request.form.get('nguoi_phu_trach', t[7] if len(t) > 7 else ''),
            request.form.get('nguoi_theo_doi', t[8] if len(t) > 8 else ''),
            request.form.get('trang_thai', t[9] if len(t) > 9 else '🟢 Đang thực hiện'),
            t[10] if len(t) > 10 else '',
            t[11] if len(t) > 11 else '',
        ]
        update_sheet(TASK_SHEET_ID, f'A{idx+1}:L{idx+1}', new)
        flash('Đã cập nhật nhiệm vụ!', 'success')
        return redirect('/tasks')
    return render_template('task_form.html', title='Sửa nhiệm vụ', icon='bi-pencil', is_edit=True, task=t)


@app.route('/tasks/status/<int:idx>', methods=['GET', 'POST'])
@login_required
def task_status(idx):
    rows = read_sheet(TASK_SHEET_ID, 'A:L')
    if idx >= len(rows):
        flash('Không tìm thấy nhiệm vụ', 'danger')
        return redirect('/tasks')
    t = rows[idx]
    if request.method == 'POST':
        new_status = request.form.get('trang_thai', '🟢 Đang thực hiện')
        update_sheet(TASK_SHEET_ID, f'J{idx+1}', [[new_status]])
        flash(f'Đã cập nhật trạng thái: {new_status}', 'success')
        return redirect('/tasks')
    return render_template('task_status.html', task=t, idx=idx)


@app.route('/tasks/delete/<int:idx>')
@login_required
def task_delete(idx):
    rows = read_sheet(TASK_SHEET_ID, 'A:L')
    if idx < len(rows):
        delete_row(TASK_SHEET_ID, idx + 1)
        flash('Đã xoá nhiệm vụ!', 'success')
    return redirect('/tasks')

# ===================== SCHEDULE =====================

@app.route('/schedule')
@login_required
def schedule():
    rows = read_sheet(LICH_SHEET_ID, 'A:I')
    return render_template('schedule.html', active='schedule', schedule=rows[2:] if len(rows) > 2 else [])


@app.route('/schedule/add', methods=['GET', 'POST'])
@login_required
def schedule_add():
    if request.method == 'POST':
        row = [
            request.form.get('ngay', ''),
            request.form.get('gio', ''),
            request.form.get('noi_dung', ''),
            request.form.get('chu_tri', ''),
            request.form.get('co_quan', ''),
            request.form.get('thanh_phan', ''),
            request.form.get('cb_vp', ''),
            request.form.get('dia_diem', ''),
            '',
        ]
        append_sheet(LICH_SHEET_ID, 'A:I', row)
        flash('Đã thêm sự kiện!', 'success')
        return redirect('/schedule')
    return render_template('schedule_form.html', title='Thêm sự kiện', icon='bi-plus-circle', is_edit=False, s=None)


@app.route('/schedule/edit/<int:idx>', methods=['GET', 'POST'])
@login_required
def schedule_edit(idx):
    rows = read_sheet(LICH_SHEET_ID, 'A:I')
    data_idx = idx + 1
    if data_idx >= len(rows):
        flash('Không tìm thấy sự kiện', 'danger')
        return redirect('/schedule')
    e = rows[data_idx] if data_idx < len(rows) else []
    if request.method == 'POST':
        new = [
            request.form.get('ngay', e[0] if len(e) > 0 else ''),
            request.form.get('gio', e[1] if len(e) > 1 else ''),
            request.form.get('noi_dung', e[2] if len(e) > 2 else ''),
            request.form.get('chu_tri', e[3] if len(e) > 3 else ''),
            request.form.get('co_quan', e[4] if len(e) > 4 else ''),
            request.form.get('thanh_phan', e[5] if len(e) > 5 else ''),
            request.form.get('cb_vp', e[6] if len(e) > 6 else ''),
            request.form.get('dia_diem', e[7] if len(e) > 7 else ''),
            '',
        ]
        update_sheet(LICH_SHEET_ID, f'A{data_idx+1}:I{data_idx+1}', new)
        flash('Đã cập nhật sự kiện!', 'success')
        return redirect('/schedule')
    return render_template('schedule_form.html', title='Sửa sự kiện', icon='bi-pencil', is_edit=True, s=e)


@app.route('/schedule/delete/<int:idx>')
@login_required
def schedule_delete(idx):
    rows = read_sheet(LICH_SHEET_ID, 'A:I')
    data_idx = idx + 1
    if data_idx < len(rows):
        delete_row(LICH_SHEET_ID, data_idx + 1)
        flash('Đã xoá sự kiện!', 'success')
    return redirect('/schedule')

# ===================== XỬ LÝ VĂN BẢN =====================

@app.route('/doc', methods=['GET', 'POST'])
@login_required
def doc():
    result = None
    if request.method == 'POST':
        f = request.files.get('file')
        ext = os.path.splitext(f.filename)[1].lower() if f else ''
        if f and ext in ('.pdf', '.docx', '.doc', '.txt'):
            path = os.path.join(SCRIPTS, f.filename)
            f.save(path)
            flash(f'Đã upload: {f.filename}', 'success')

            def _extract(fp):
                e = os.path.splitext(fp)[1].lower()
                if e == '.pdf':
                    import pdfplumber
                    with pdfplumber.open(fp) as p:
                        return '\n'.join(page.extract_text() or '' for page in p.pages)
                elif e in ('.docx', '.doc'):
                    from docx import Document
                    doc = Document(fp)
                    return '\n'.join(p.text for p in doc.paragraphs)
                else:
                    with open(fp, 'r', encoding='utf-8') as fh:
                        return fh.read()

            def _parse(text, filename):
                from agents.agent_calendar import parse_time, parse_location
                import re
                lines = text.split('\n')
                info = {
                    'file': filename, 'text': text, 'content': '',
                    'location': '', 'time_str': '', 'chu_tri': '',
                    'thanh_phan': [], 'so_hieu': os.path.splitext(filename)[0],
                    **parse_time(text),
                }
                for i, line in enumerate(lines):
                    ls = line.strip()
                    m = re.match(r'^[-•]?\s*Nội dung\s*:\s*(.+)', ls, re.I)
                    if m: info['content'] = m.group(1).strip(); break
                    if ls.lower().startswith('về nội dung') or ls.lower().startswith('về việc'):
                        parts = [ls]
                        for j in range(i+1, min(i+5, len(lines))):
                            nxt = lines[j].strip()
                            if not nxt or any(kw in nxt.lower() for kw in ['kính gửi', 'kính mời', 'số:']): break
                            parts.append(nxt)
                        info['content'] = ' '.join(parts); break
                if not info['content']:
                    for line in lines[:15]:
                        ls = line.strip()
                        if ls and len(ls) > 20 and not any(kw in ls.upper() for kw in ['ỦY BAN', 'CỘNG HÒA', 'ĐỘC LẬP', 'SỐ:', 'GIẤY MỜI', 'KÍNH GỬI']):
                            info['content'] = ls; break
                info['location'] = parse_location(text)
                for line in lines:
                    ls = line.strip()
                    if 'chủ trì' in ls.lower():
                        info['chu_tri'] = re.sub(r'^[-•\s]*(Chủ trì|chủ trì)\s*[:\-]?\s*', '', ls).strip().rstrip('.,;')
                        break
                in_kg = False
                for line in lines:
                    ls = line.strip()
                    if 'kính gửi' in ls.lower(): in_kg = True; continue
                    if in_kg:
                        if ls.startswith('-') or ls.startswith('•'):
                            t = ls.lstrip('-• ').strip().rstrip('.,;')
                            if t: info['thanh_phan'].append(t)
                        elif info['thanh_phan']:
                            if len(ls) < 20 and ls and not any(kw in ls.lower() for kw in ['kính mời', 'nội dung', 'thời gian']):
                                info['thanh_phan'][-1] += ' ' + ls.rstrip('.,;')
                            else: break
                return info

            text = _extract(path)
            info = _parse(text, f.filename)
            result = info
            from agents.agent_calendar import load_services, to_calendar_event, to_sheets_row
            try:
                cal, sh = load_services()
                ev = to_calendar_event(info)
                if ev:
                    cal.events().insert(calendarId='primary', body=ev).execute()
                if LICH_SHEET_ID:
                    append_sheet(LICH_SHEET_ID, 'A:I', to_sheets_row(info))
                if TASK_SHEET_ID:
                    now = date.today()
                    tr = ["", info["content"], info["so_hieu"] or os.path.splitext(f.filename)[0],
                          now.strftime("%d/%m/%Y"), "", info["chu_tri"],
                          "; ".join(info["thanh_phan"][:3]) if info["thanh_phan"] else "",
                          "", "", "🟢 Đang thực hiện", "", ""]
                    append_sheet(TASK_SHEET_ID, 'A:L', tr)
                flash('✅ Đã xử lý: Calendar + Sheets đã cập nhật!', 'success')
            except Exception as e:
                flash(f'⚠️ Lỗi xử lý: {e}', 'warning')

    pdf_files = []
    for ext in ('*.pdf', '*.docx', '*.doc', '*.txt'):
        for p in sorted(glob.glob(os.path.join(SCRIPTS, ext))):
            size = os.path.getsize(p)
            if size < 1024:
                sz = f'{size} B'
            elif size < 1024*1024:
                sz = f'{size//1024} KB'
            else:
                sz = f'{size/(1024*1024):.1f} MB'
            pdf_files.append({
                'name': os.path.basename(p),
                'size': sz,
                'date': datetime.fromtimestamp(os.path.getmtime(p)).strftime('%d/%m/%Y %H:%M')
            })
    return render_template('doc.html', active='doc', files=pdf_files, result=result)


@app.route('/pdf', methods=['GET', 'POST'])
@login_required
def pdf_redirect():
    return redirect('/doc')


@app.route('/doc/process/<int:idx>')
@login_required
def doc_process(idx):
    all_files = []
    for ext in ('*.pdf', '*.docx', '*.doc', '*.txt'):
        all_files.extend(sorted(glob.glob(os.path.join(SCRIPTS, ext))))
    if idx < len(all_files):
        path = all_files[idx]
        def extract_text(fp):
            e = os.path.splitext(fp)[1].lower()
            if e == '.pdf':
                import pdfplumber
                with pdfplumber.open(fp) as p:
                    return '\n'.join(page.extract_text() or '' for page in p.pages)
            elif e in ('.docx', '.doc'):
                from docx import Document
                doc = Document(fp)
                return '\n'.join(p.text for p in doc.paragraphs)
            else:
                with open(fp, 'r', encoding='utf-8') as fh:
                    return fh.read()
        def parse_text_to_info(text, filename):
            from agents.agent_calendar import parse_time, parse_location
            import re
            lines = text.split('\n')
            info = {
                'file': filename, 'text': text, 'content': '',
                'location': '', 'time_str': '', 'chu_tri': '',
                'thanh_phan': [], 'so_hieu': os.path.splitext(filename)[0],
                **parse_time(text),
            }
            for i, line in enumerate(lines):
                ls = line.strip()
                m = re.match(r'^[-•]?\s*Nội dung\s*:\s*(.+)', ls, re.I)
                if m: info['content'] = m.group(1).strip(); break
                if ls.lower().startswith('về nội dung') or ls.lower().startswith('về việc'):
                    parts = [ls]
                    for j in range(i+1, min(i+5, len(lines))):
                        nxt = lines[j].strip()
                        if not nxt or any(kw in nxt.lower() for kw in ['kính gửi', 'kính mời', 'số:']): break
                        parts.append(nxt)
                    info['content'] = ' '.join(parts); break
            if not info['content']:
                for line in lines[:15]:
                    ls = line.strip()
                    if ls and len(ls) > 20 and not any(kw in ls.upper() for kw in ['ỦY BAN', 'CỘNG HÒA', 'ĐỘC LẬP', 'SỐ:', 'GIẤY MỜI', 'KÍNH GỬI']):
                        info['content'] = ls; break
            info['location'] = parse_location(text)
            for line in lines:
                ls = line.strip()
                if 'chủ trì' in ls.lower():
                    info['chu_tri'] = re.sub(r'^[-•\s]*(Chủ trì|chủ trì)\s*[:\-]?\s*', '', ls).strip().rstrip('.,;')
                    break
            in_kg = False
            for line in lines:
                ls = line.strip()
                if 'kính gửi' in ls.lower(): in_kg = True; continue
                if in_kg:
                    if ls.startswith('-') or ls.startswith('•'):
                        t = ls.lstrip('-• ').strip().rstrip('.,;')
                        if t: info['thanh_phan'].append(t)
                    elif info['thanh_phan']:
                        if len(ls) < 20 and ls and not any(kw in ls.lower() for kw in ['kính mời', 'nội dung', 'thời gian']):
                            info['thanh_phan'][-1] += ' ' + ls.rstrip('.,;')
                        else: break
            return info
        text = extract_text(path)
        info = parse_text_to_info(text, os.path.basename(path))
        try:
            cal, sh = load_services()
            ev = to_calendar_event(info)
            if ev:
                cal.events().insert(calendarId='primary', body=ev).execute()
            if LICH_SHEET_ID:
                append_sheet(LICH_SHEET_ID, 'A:I', to_sheets_row(info))
            if TASK_SHEET_ID:
                now = date.today()
                tp_str = "; ".join(info["thanh_phan"][:3]) if info["thanh_phan"] else ""
                dvcb = info.get("don_vi_chuan_bi", "")
                phoi_hop = "; ".join(filter(None, [tp_str, dvcb]))
                tr = ["", info["content"], info["so_hieu"] or os.path.splitext(os.path.basename(path))[0],
                      now.strftime("%d/%m/%Y"), "", info["chu_tri"], phoi_hop,
                      "", "", "🟢 Đang thực hiện", "", ""]
                append_sheet(TASK_SHEET_ID, 'A:L', tr)
            flash(f'✅ Đã xử lý: {os.path.basename(path)}', 'success')
        except Exception as e:
            flash(f'⚠️ Lỗi: {e}', 'danger')
    return redirect('/doc')


@app.route('/pdf/process/<int:idx>')
@login_required
def pdf_process(idx):
    all_files = []
    for ext in ('*.pdf', '*.docx', '*.doc', '*.txt'):
        all_files.extend(sorted(glob.glob(os.path.join(SCRIPTS, ext))))
    if idx < len(all_files):
        path = all_files[idx]
        from agents.agent_calendar import parse_pdf, load_services, to_calendar_event, to_sheets_row
        def extract_text(fp):
            e = os.path.splitext(fp)[1].lower()
            if e == '.pdf':
                import pdfplumber
                with pdfplumber.open(fp) as p:
                    return '\n'.join(page.extract_text() or '' for page in p.pages)
            elif e in ('.docx', '.doc'):
                from docx import Document
                doc = Document(fp)
                return '\n'.join(p.text for p in doc.paragraphs)
            else:
                with open(fp, 'r', encoding='utf-8') as fh:
                    return fh.read()

        def parse_text_to_info(text, filename):
            from agents.agent_calendar import parse_time, parse_location
            import re
            lines = text.split('\n')
            info = {
                'file': filename, 'text': text, 'content': '',
                'location': '', 'time_str': '', 'chu_tri': '',
                'thanh_phan': [], 'so_hieu': os.path.splitext(filename)[0],
                **parse_time(text),
            }
            for i, line in enumerate(lines):
                ls = line.strip()
                m = re.match(r'^[-•]?\s*Nội dung\s*:\s*(.+)', ls, re.I)
                if m:
                    info['content'] = m.group(1).strip(); break
                if ls.lower().startswith('về nội dung') or ls.lower().startswith('về việc'):
                    parts = [ls]
                    for j in range(i+1, min(i+5, len(lines))):
                        nxt = lines[j].strip()
                        if not nxt or any(kw in nxt.lower() for kw in ['kính gửi', 'kính mời', 'số:']): break
                        parts.append(nxt)
                    info['content'] = ' '.join(parts); break
            if not info['content']:
                for line in lines[:15]:
                    ls = line.strip()
                    if ls and len(ls) > 20 and not any(kw in ls.upper() for kw in ['ỦY BAN', 'CỘNG HÒA', 'ĐỘC LẬP', 'SỐ:', 'GIẤY MỜI', 'KÍNH GỬI']):
                        info['content'] = ls; break
            info['location'] = parse_location(text)
            for line in lines:
                ls = line.strip()
                if 'chủ trì' in ls.lower():
                    info['chu_tri'] = re.sub(r'^[-•\s]*(Chủ trì|chủ trì)\s*[:\-]?\s*', '', ls).strip().rstrip('.,;')
                    break
            info['thanh_phan'] = []
            in_kg = False
            for line in lines:
                ls = line.strip()
                if 'kính gửi' in ls.lower(): in_kg = True; continue
                if in_kg:
                    if ls.startswith('-') or ls.startswith('•'):
                        t = ls.lstrip('-• ').strip().rstrip('.,;')
                        if t: info['thanh_phan'].append(t)
                    elif info['thanh_phan']:
                        if len(ls) < 20 and ls and not any(kw in ls.lower() for kw in ['kính mời', 'nội dung', 'thời gian']):
                            info['thanh_phan'][-1] += ' ' + ls.rstrip('.,;')
                        else: break
            return info

        text = extract_text(path)
        info = parse_text_to_info(text, os.path.basename(path))
        try:
            cal, sh = load_services()
            ev = to_calendar_event(info)
            if ev:
                cal.events().insert(calendarId='primary', body=ev).execute()
            if LICH_SHEET_ID:
                append_sheet(LICH_SHEET_ID, 'A:I', to_sheets_row(info))
            if TASK_SHEET_ID:
                now = date.today()
                tp_str = "; ".join(info["thanh_phan"][:3]) if info["thanh_phan"] else ""
                dvcb = info.get("don_vi_chuan_bi", "")
                phoi_hop = "; ".join(filter(None, [tp_str, dvcb]))
                tr = ["", info["content"], info["so_hieu"] or os.path.splitext(os.path.basename(path))[0],
                      now.strftime("%d/%m/%Y"), "", info["chu_tri"], phoi_hop,
                      "", "", "🟢 Đang thực hiện", "", ""]
                append_sheet(TASK_SHEET_ID, 'A:L', tr)
            flash(f'✅ Đã xử lý: {os.path.basename(path)}', 'success')
        except Exception as e:
            flash(f'⚠️ Lỗi: {e}', 'danger')
    return redirect('/doc')

# ===================== CALENDAR =====================

@app.route('/calendar')
@login_required
def calendar():
    try:
        cal = cal_svc()
        now = datetime.now(timezone.utc)
        past = (now - timedelta(days=14)).isoformat()
        future = (now + timedelta(days=60)).isoformat()
        items = cal.events().list(
            calendarId=CALENDAR_ID, timeMin=past, timeMax=future,
            maxResults=50, singleEvents=True,
            orderBy='startTime'
        ).execute().get('items', [])
        seen = set()
        events = []
        for ev in items:
            start = ev['start'].get('dateTime', ev['start'].get('date', ''))[:16]
            summary = ev.get('summary', '(không tiêu đề)')[:60]
            key = start + summary
            if key in seen:
                continue
            seen.add(key)
            is_all_day = 'date' in ev['start'] and 'dateTime' not in ev['start']
            events.append({
                'id': ev['id'],
                'summary': summary,
                'start': start,
                'end': ev['end'].get('dateTime', ev['end'].get('date', ''))[:16],
                'location': ev.get('location', '')[:40],
                'description': ev.get('description', '')[:100],
                'is_all_day': is_all_day,
            })
    except Exception as e:
        flash(f'Lỗi Calendar: {e}', 'danger')
        events = []
    return render_template('calendar.html', active='calendar', events=events)


@app.route('/calendar/add', methods=['GET', 'POST'])
@login_required
def calendar_add():
    if request.method == 'POST':
        ngay = request.form.get('ngay', '').strip()
        gio_bat_dau = request.form.get('gio_bd', '08:00').strip()
        gio_ket_thuc = request.form.get('gio_kt', '09:00').strip()
        is_all_day = request.form.get('all_day') == 'on'

        if not ngay:
            flash('Vui lòng nhập ngày', 'danger')
            return redirect('/calendar/add')

        try:
            d = ngay.split('/')
            ngay_iso = f'{d[2]}-{d[1]}-{d[0]}'
            if is_all_day:
                start = {'date': ngay_iso}
                end = {'date': ngay_iso}
            else:
                start = {'dateTime': f'{ngay_iso}T{gio_bat_dau}:00', 'timeZone': 'Asia/Ho_Chi_Minh'}
                end = {'dateTime': f'{ngay_iso}T{gio_ket_thuc}:00', 'timeZone': 'Asia/Ho_Chi_Minh'}

            body = {
                'summary': request.form.get('summary', ''),
                'start': start,
                'end': end,
                'location': request.form.get('location', ''),
                'description': request.form.get('description', ''),
            }
            cal = cal_svc()
            cal.events().insert(calendarId=CALENDAR_ID, body=body).execute()
            flash('✅ Đã thêm sự kiện vào Calendar!', 'success')
        except Exception as e:
            flash(f'Lỗi: {e}', 'danger')
        return redirect('/calendar')

    return render_template('cal_form.html', active='calendar', is_edit=False, ev=None, ngay='', gio_bd='08:00', gio_kt='09:00', is_all_day=False)


@app.route('/calendar/edit/<event_id>', methods=['GET', 'POST'])
@login_required
def calendar_edit(event_id):
    cal = cal_svc()
    try:
        ev = cal.events().get(calendarId=CALENDAR_ID, eventId=event_id).execute()
    except Exception as e:
        flash(f'Không tìm thấy sự kiện: {e}', 'danger')
        return redirect('/calendar')

    if request.method == 'POST':
        ngay = request.form.get('ngay', '').strip()
        gio_bd = request.form.get('gio_bd', '08:00').strip()
        gio_kt = request.form.get('gio_kt', '09:00').strip()
        is_all_day = request.form.get('all_day') == 'on'

        if ngay:
            try:
                d = ngay.split('/')
                ngay_iso = f'{d[2]}-{d[1]}-{d[0]}'
                if is_all_day:
                    ev['start'] = {'date': ngay_iso}
                    ev['end'] = {'date': ngay_iso}
                else:
                    ev['start'] = {'dateTime': f'{ngay_iso}T{gio_bd}:00', 'timeZone': 'Asia/Ho_Chi_Minh'}
                    ev['end'] = {'dateTime': f'{ngay_iso}T{gio_kt}:00', 'timeZone': 'Asia/Ho_Chi_Minh'}
            except Exception:
                pass

        ev['summary'] = request.form.get('summary', ev.get('summary', ''))
        ev['location'] = request.form.get('location', '')
        ev['description'] = request.form.get('description', '')

        try:
            cal.events().update(calendarId=CALENDAR_ID, eventId=event_id, body=ev).execute()
            flash('✅ Đã cập nhật sự kiện!', 'success')
        except Exception as e:
            flash(f'Lỗi: {e}', 'danger')
        return redirect('/calendar')

    is_all_day = 'date' in ev['start'] and 'dateTime' not in ev['start']
    dt = ev['start'].get('dateTime', ev['start'].get('date', ''))[:16]
    ngay = gio_bd = gio_kt = ''
    if dt:
        parts = dt.split('T')
        ngay_p = parts[0].split('-')
        if len(ngay_p) == 3:
            ngay = f'{ngay_p[2]}/{ngay_p[1]}/{ngay_p[0]}'
        if len(parts) > 1:
            gio_bd = parts[1][:5]
    if not is_all_day and 'dateTime' in ev.get('end', {}):
        end_dt = ev['end']['dateTime'][:16]
        if 'T' in end_dt:
            gio_kt = end_dt.split('T')[1][:5]
    return render_template('cal_form.html', active='calendar', is_edit=True,
                           ev=ev, ngay=ngay, gio_bd=gio_bd, gio_kt=gio_kt, is_all_day=is_all_day)


@app.route('/calendar/delete/<event_id>')
@login_required
def calendar_delete(event_id):
    try:
        cal = cal_svc()
        cal.events().delete(calendarId=CALENDAR_ID, eventId=event_id).execute()
        flash('✅ Đã xoá sự kiện!', 'success')
    except Exception as e:
        flash(f'Lỗi: {e}', 'danger')
    return redirect('/calendar')

# ===================== SYNC & RUN =====================

@app.route('/sync')
@login_required
def sync():
    return render_template('sync.html', active='sync')


@app.route('/sync/sheets-to-web')
@login_required
def sync_sheets_to_web():
    flash('✅ Dữ liệu đã được đồng bộ từ Google Sheets!', 'success')
    return redirect('/')


@app.route('/sync/web-to-sheets')
@login_required
def sync_web_to_sheets():
    rows = read_sheet(TASK_SHEET_ID, 'A:L')
    updated = 0
    for i, row in enumerate(rows[1:], 2):
        if len(row) > 9 and 'Hoàn thành' in row[9]:
            continue
        deadline = row[4] if len(row) > 4 else ""
        if deadline:
            try:
                p = deadline.strip().split('/')
                d = date(int(p[2]), int(p[1]), int(p[0]))
                diff = (d - date.today()).days
                if diff < 0:
                    ns = "🔴 Quá hạn"
                elif diff <= 3:
                    ns = "🟡 Sắp đến hạn"
                else:
                    ns = "🟢 Đang thực hiện"
                update_sheet(TASK_SHEET_ID, f'J{i}', [[ns]])
                updated += 1
            except:
                pass
    flash(f'✅ Đã đồng bộ {updated} nhiệm vụ lên Google Sheets', 'success')
    return redirect('/')


@app.route('/pdf/upload', methods=['POST'])
@login_required
def pdf_upload():
    f = request.files.get('file')
    if not f:
        flash('Vui lòng chọn file', 'danger')
        return redirect('/doc')
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in ('.pdf', '.docx', '.doc', '.txt'):
        flash('Chỉ hỗ trợ PDF, DOCX, DOC, TXT', 'danger')
        return redirect('/doc')
    path = os.path.join(SCRIPTS, f.filename)
    f.save(path)

    from agents.agent_calendar import load_services, to_calendar_event, to_sheets_row

    def extract_text(filepath):
        e = os.path.splitext(filepath)[1].lower()
        if e == '.pdf':
            import pdfplumber
            with pdfplumber.open(filepath) as p:
                return '\n'.join(page.extract_text() or '' for page in p.pages)
        elif e in ('.docx', '.doc'):
            from docx import Document
            doc = Document(filepath)
            return '\n'.join(p.text for p in doc.paragraphs)
        else:
            with open(filepath, 'r', encoding='utf-8') as fh:
                return fh.read()

    def parse_text_to_info(text, filename):
        from agents.agent_calendar import parse_time, parse_location
        import re
        lines = text.split('\n')
        info = {
            'file': filename, 'text': text, 'content': '',
            'location': '', 'time_str': '', 'chu_tri': '',
            'thanh_phan': [], 'so_hieu': os.path.splitext(filename)[0],
            **parse_time(text),
        }
        for i, line in enumerate(lines):
            ls = line.strip()
            m = re.match(r'^[-•]?\s*Nội dung\s*:\s*(.+)', ls, re.I)
            if m:
                info['content'] = m.group(1).strip(); break
            if ls.lower().startswith('về nội dung') or ls.lower().startswith('về việc'):
                parts = [ls]
                for j in range(i+1, min(i+5, len(lines))):
                    nxt = lines[j].strip()
                    if not nxt or any(kw in nxt.lower() for kw in ['kính gửi', 'kính mời', 'số:']): break
                    parts.append(nxt)
                info['content'] = ' '.join(parts); break
        if not info['content']:
            for line in lines[:15]:
                ls = line.strip()
                if ls and len(ls) > 20 and not any(kw in ls.upper() for kw in ['ỦY BAN', 'CỘNG HÒA', 'ĐỘC LẬP', 'SỐ:', 'GIẤY MỜI', 'KÍNH GỬI']):
                    info['content'] = ls; break
        info['location'] = parse_location(text)
        for line in lines:
            ls = line.strip()
            if 'chủ trì' in ls.lower():
                info['chu_tri'] = re.sub(r'^[-•\s]*(Chủ trì|chủ trì)\s*[:\-]?\s*', '', ls).strip().rstrip('.,;')
                break
        info['thanh_phan'] = []
        in_kg = False
        for line in lines:
            ls = line.strip()
            if 'kính gửi' in ls.lower(): in_kg = True; continue
            if in_kg:
                if ls.startswith('-') or ls.startswith('•'):
                    t = ls.lstrip('-• ').strip().rstrip('.,;')
                    if t: info['thanh_phan'].append(t)
                elif info['thanh_phan']:
                    if len(ls) < 20 and ls and not any(kw in ls.lower() for kw in ['kính mời', 'nội dung', 'thời gian']):
                        info['thanh_phan'][-1] += ' ' + ls.rstrip('.,;')
                    else: break
        return info

    try:
        text = extract_text(path)
        info = parse_text_to_info(text, f.filename)
        if request.form.get('to_calendar'):
            cal, _ = load_services()
            ev = to_calendar_event(info)
            if ev:
                cal.events().insert(calendarId='primary', body=ev).execute()
        if request.form.get('to_sheets'):
            if LICH_SHEET_ID:
                append_sheet(LICH_SHEET_ID, 'A:I', to_sheets_row(info))
            if TASK_SHEET_ID:
                tp_str = "; ".join(info["thanh_phan"][:3]) if info["thanh_phan"] else ""
                dvcb = info.get("don_vi_chuan_bi", "")
                phoi_hop = "; ".join(filter(None, [tp_str, dvcb]))
                tr = ["", info["content"], info["so_hieu"] or os.path.splitext(f.filename)[0],
                      "", "", info["chu_tri"], phoi_hop,
                      "", "", "🟢 Đang thực hiện", "", ""]
                append_sheet(TASK_SHEET_ID, 'A:L', tr)
        flash(f'✅ Đã xử lý: {f.filename}', 'success')
    except Exception as e:
        flash(f'⚠️ Lỗi xử lý: {e}', 'danger')
    return redirect('/doc')


# ===================== SETTINGS =====================

ENV_PATH = os.path.join(SCRIPTS, '.env')


def read_env():
    cfg = {}
    for line in open(ENV_PATH, encoding='utf-8'):
        line = line.strip()
        if line and not line.startswith('#') and '=' in line:
            k, v = line.split('=', 1)
            cfg[k.strip()] = v.strip()
    return cfg


def write_env(cfg):
    lines = []
    for line in open(ENV_PATH, encoding='utf-8'):
        stripped = line.strip()
        if stripped and not stripped.startswith('#') and '=' in stripped:
            k = stripped.split('=', 1)[0].strip()
            if k in cfg:
                line = f'{k}={cfg[k]}\n'
                del cfg[k]
        lines.append(line)
    for k, v in cfg.items():
        lines.append(f'{k}={v}\n')
    open(ENV_PATH, 'w', encoding='utf-8').writelines(lines)


@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    if request.method == 'POST':
        cfg = {
            'GOOGLE_CALENDAR_ID': request.form.get('calendar_id', 'primary'),
            'GOOGLE_SHEET_TASK_ID': request.form.get('task_sheet_id', ''),
            'GOOGLE_SHEET_LICH_ID': request.form.get('lich_sheet_id', ''),
            'NOTIFICATION_EMAIL': request.form.get('email', ''),
        }
        write_env(cfg)
        global CALENDAR_ID, TASK_SHEET_ID, LICH_SHEET_ID
        CALENDAR_ID = cfg['GOOGLE_CALENDAR_ID']
        TASK_SHEET_ID = cfg['GOOGLE_SHEET_TASK_ID']
        LICH_SHEET_ID = cfg['GOOGLE_SHEET_LICH_ID']
        load_dotenv(dotenv_path=ENV_PATH, override=True)
        flash('✅ Đã lưu cấu hình!', 'success')
        return redirect('/settings')

    cfg = read_env()
    oauth_ok = os.path.exists(token_path)
    return render_template('settings.html', active='settings',
                           config=cfg, oauth_ok=oauth_ok)


@app.route('/settings/re-auth')
@login_required
def settings_re_auth():
    import subprocess
    subprocess.Popen([sys.executable, os.path.join(SCRIPTS, 'auth_init.py')],
                     creationflags=subprocess.CREATE_NEW_CONSOLE)
    flash('🔄 Đang mở cửa sổ xác thực Google OAuth...', 'info')
    return redirect('/settings')


# ===================== AGENT 3: KẾT LUẬN =====================

@app.route('/ket-luan', methods=['GET', 'POST'])
@login_required
def ket_luan():
    result = None
    files = []
    out_dir = os.path.join(BASE, "output", "ket_luan")
    os.makedirs(out_dir, exist_ok=True)

    if request.method == 'POST':
        f = request.files.get('file')
        if f and f.filename.lower().endswith(('.pdf', '.docx', '.txt')):
            path = os.path.join(SCRIPTS, f.filename)
            f.save(path)

            from agents.agent_ket_luan import run as run_agent3
            try:
                out_path = run_agent3(input_path=path)
                result = {
                    'file': os.path.basename(path),
                    'output': os.path.basename(out_path) if out_path else '',
                }
                flash(f'✅ Đã tạo Thông báo kết luận từ: {f.filename}', 'success')
            except Exception as e:
                flash(f'⚠️ Lỗi xử lý: {e}', 'danger')
        else:
            flash('Vui lòng chọn file PDF/DOCX/TXT', 'danger')

    for p in sorted(glob.glob(os.path.join(out_dir, '*.*')), key=os.path.getmtime, reverse=True):
        size = os.path.getsize(p)
        if size < 1024:
            sz = f'{size} B'
        elif size < 1024 * 1024:
            sz = f'{size // 1024} KB'
        else:
            sz = f'{size / (1024 * 1024):.1f} MB'
        files.append({
            'name': os.path.basename(p),
            'size': sz,
            'date': datetime.fromtimestamp(os.path.getmtime(p)).strftime('%d/%m/%Y %H:%M'),
            'path': p,
        })

    return render_template('ket_luan.html', active='ket_luan', files=files, result=result)


@app.route('/ket-luan/download/<filename>')
@login_required
def ket_luan_download(filename):
    out_dir = os.path.join(BASE, "output", "ket_luan")
    return send_from_directory(out_dir, filename, as_attachment=True)


APP_SCRIPTS_DIR = SCRIPTS


@app.route('/run-all')
@login_required
def run_all():
    import subprocess
    subprocess.Popen([sys.executable, os.path.join(SCRIPTS, 'run_all.py')],
                     creationflags=subprocess.CREATE_NEW_CONSOLE)
    flash('🚀 Đang chạy tất cả Agent (cửa sổ mới)...', 'info')
    return redirect('/')


if __name__ == '__main__':
    port = 5000
    print(f"🌐 UBND Bà Rịa - Web App: http://localhost:{port}")
    print(f"   Nhấn Ctrl+C để dừng")
    app.run(debug=True, port=port, host='0.0.0.0')
